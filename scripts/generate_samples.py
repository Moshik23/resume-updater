"""Dev-only helper: generates samples/sample_resume.docx and sample_resume.pdf
for Phase 0 local testing. Not part of the app runtime.

Deliberately includes a two-column header table (to exercise the docx
ingest's table-walking path) and omits Kubernetes / a cloud certification,
so the JD gap-analysis has something real to flag.
"""

import os

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

SAMPLES_DIR = os.path.join(os.path.dirname(__file__), "..", "samples")


def build_docx(path: str) -> None:
    doc = Document()

    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].paragraphs[0].add_run("Jordan Rivera").bold = True
    table.rows[0].cells[1].paragraphs[0].add_run("jordan.rivera@example.com | (555) 123-4567")

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(
        "Cloud engineer with 3 years of experience building and operating Azure infrastructure."
    )

    doc.add_heading("Experience", level=2)
    doc.add_paragraph("Cloud Engineer, Example Corp (2022-Present)", style="Heading 3")
    doc.add_paragraph("Managed Azure virtual machines and storage accounts for production workloads.", style="List Bullet")
    doc.add_paragraph("Wrote Terraform modules to provision networking and compute resources.", style="List Bullet")
    doc.add_paragraph("Built CI/CD pipelines in Azure DevOps for automated deployments.", style="List Bullet")

    doc.add_heading("Certifications and Tech Skills", level=2)
    doc.add_paragraph("Cloud Platforms: Microsoft Azure (VMs, Storage, Networking)", style="List Bullet")
    doc.add_paragraph("Infrastructure as Code: Terraform", style="List Bullet")
    doc.add_paragraph("CI/CD: Azure DevOps", style="List Bullet")
    doc.add_paragraph("Scripting: Python, Bash", style="List Bullet")
    doc.add_paragraph("Containerisation: Docker", style="List Bullet")

    doc.save(path)


def build_pdf(path: str) -> None:
    c = canvas.Canvas(path, pagesize=LETTER)
    lines = [
        "JORDAN RIVERA",
        "jordan.rivera@example.com | (555) 123-4567",
        "",
        "SUMMARY",
        "Cloud engineer with 3 years of experience building and operating Azure infrastructure.",
        "",
        "EXPERIENCE",
        "Cloud Engineer, Example Corp (2022-Present)",
        "* Managed Azure virtual machines and storage accounts for production workloads.",
        "* Wrote Terraform modules to provision networking and compute resources.",
        "* Built CI/CD pipelines in Azure DevOps for automated deployments.",
        "",
        "CERTIFICATIONS AND TECH SKILLS",
        "* Cloud Platforms: Microsoft Azure (VMs, Storage, Networking)",
        "* Infrastructure as Code: Terraform",
        "* CI/CD: Azure DevOps",
        "* Scripting: Python, Bash",
        "* Containerisation: Docker",
    ]
    y = 740
    for line in lines:
        c.drawString(72, y, line)
        y -= 18
    c.save()


if __name__ == "__main__":
    os.makedirs(SAMPLES_DIR, exist_ok=True)
    build_docx(os.path.join(SAMPLES_DIR, "sample_resume.docx"))
    build_pdf(os.path.join(SAMPLES_DIR, "sample_resume.pdf"))
    print("Generated sample_resume.docx and sample_resume.pdf")
