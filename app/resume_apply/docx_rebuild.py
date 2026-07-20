"""Renders a clean, ATS-friendly .docx from extracted (and edited) content
blocks -- used only for the PDF input path, where there's no original docx
paragraph/run structure to preserve. This is an explicit re-render, not a
layout clone of the source PDF; that trade-off is communicated to the user
in the UI and README rather than silently attempted and falling short.
"""

import re

from docx import Document

from app.models import ContentBlock, SuggestedEdit

_BULLET_PREFIX = re.compile(r"^[•\-*●○]\s*")


def _guess_kind(text: str, is_first: bool) -> str:
    if is_first:
        return "title"
    if _BULLET_PREFIX.match(text):
        return "bullet"
    words = text.split()
    if text.isupper() and 1 <= len(words) <= 6:
        return "heading"
    return "body"


def _apply_edits_to_blocks(
    blocks: list[ContentBlock], edits: list[SuggestedEdit]
) -> tuple[list[ContentBlock], list[str]]:
    result = list(blocks)
    position = {b.id: i for i, b in enumerate(result)}
    failed: list[str] = []
    next_ordinal = 0

    def fresh_id() -> str:
        nonlocal next_ordinal
        next_ordinal += 1
        return f"new{next_ordinal}"

    for edit in edits:
        idx = position.get(edit.block_id)
        if idx is None:
            failed.append(edit.edit_id)
            continue
        block = result[idx]

        if edit.type in ("replace_phrase", "insert_keyword"):
            if edit.anchor_text and edit.anchor_text in block.text:
                result[idx] = block.model_copy(
                    update={"text": block.text.replace(edit.anchor_text, edit.new_text, 1)}
                )
            else:
                failed.append(edit.edit_id)
        elif edit.type in ("append_bullet", "new_bullet_after"):
            insert_at = idx + 1
            new_block = ContentBlock(id=fresh_id(), text=f"• {edit.new_text}")
            result.insert(insert_at, new_block)
            for key, pos in position.items():
                if pos >= insert_at:
                    position[key] = pos + 1
            # Advance so a second new bullet for the same block_id lands
            # after this one instead of both landing right after the anchor.
            position[edit.block_id] = insert_at
        else:
            failed.append(edit.edit_id)

    return result, failed


def build(blocks: list[ContentBlock], edits: list[SuggestedEdit], output_path: str) -> list[str]:
    """Apply edits to the extracted blocks and render a clean docx.

    Returns the edit_ids that could not be applied.
    """
    final_blocks, failed = _apply_edits_to_blocks(blocks, edits)

    doc = Document()
    for i, block in enumerate(final_blocks):
        kind = _guess_kind(block.text, is_first=(i == 0))
        if kind == "title":
            doc.add_heading(block.text, level=0)
        elif kind == "heading":
            doc.add_heading(block.text, level=2)
        elif kind == "bullet":
            doc.add_paragraph(_BULLET_PREFIX.sub("", block.text), style="List Bullet")
        else:
            doc.add_paragraph(block.text)

    doc.save(output_path)
    return failed
