"""Extracts a flat, deterministic content-block registry from a .docx resume.

The same registry (paragraph objects keyed by block id) must be re-derivable
from the same file bytes at both the analysis step and the (possibly much
later, cold-started) apply step, so ids are assigned purely by walk order —
no random ids, no cached state.
"""

from collections.abc import Iterator

from docx import Document
from docx.text.paragraph import Paragraph

from app.models import ContentBlock


def load(path: str) -> Document:
    return Document(path)


def _iter_paragraphs_with_ids(doc: Document) -> Iterator[tuple[str, Paragraph]]:
    for i, p in enumerate(doc.paragraphs):
        yield f"p{i}", p
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, p in enumerate(cell.paragraphs):
                    yield f"t{t_idx}r{r_idx}c{c_idx}p{p_idx}", p


def get_paragraph_registry(doc: Document) -> dict[str, Paragraph]:
    """Rebuild the block_id -> Paragraph mapping for an already-opened Document."""
    return dict(_iter_paragraphs_with_ids(doc))


def extract(path: str) -> list[ContentBlock]:
    """Extract non-empty paragraphs (body + table cells) as content blocks."""
    doc = load(path)
    blocks: list[ContentBlock] = []
    for block_id, paragraph in _iter_paragraphs_with_ids(doc):
        text = paragraph.text
        if not text.strip():
            continue
        style_name = paragraph.style.name if paragraph.style is not None else None
        blocks.append(ContentBlock(id=block_id, text=text, style_name=style_name))
    return blocks
